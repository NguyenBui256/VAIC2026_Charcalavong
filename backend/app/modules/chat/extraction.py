"""Safe, bounded text extraction and attachment-context selection."""

from __future__ import annotations

import csv
import io
import re
import zipfile
from pathlib import Path

from app.core.errors import ValidationError

MAX_EXTRACTED_CHARS = 200_000
MAX_CONTEXT_CHARS = 60_000
CHUNK_CHARS = 1_500

ALLOWED_TYPES = {
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    ".txt": {"text/plain"},
    ".md": {"text/markdown", "text/plain"},
    ".csv": {"text/csv", "application/csv", "text/plain", "application/vnd.ms-excel"},
    ".xlsx": {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
}


def validate_file_type(filename: str, content_type: str, path: Path) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_TYPES or suffix in {".docm", ".xlsm"}:
        raise ValidationError("unsupported attachment type", code="unsupported_file_type")
    normalized = (content_type or "").split(";", 1)[0].strip().lower()
    if normalized not in ALLOWED_TYPES[suffix]:
        raise ValidationError("file extension and MIME type do not match", code="mime_mismatch")
    header = path.read_bytes()[:8]
    if suffix == ".pdf" and not header.startswith(b"%PDF-"):
        raise ValidationError("invalid PDF signature", code="mime_mismatch")
    if suffix in {".docx", ".xlsx"} and not header.startswith(b"PK"):
        raise ValidationError("invalid Office document signature", code="mime_mismatch")
    return suffix


def extract_text(path: Path, filename: str, content_type: str) -> str:
    suffix = validate_file_type(filename, content_type, path)
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8-sig")
    elif suffix == ".csv":
        text = _extract_csv(path)
    else:
        text = _extract_xlsx(path)
    return text[:MAX_EXTRACTED_CHARS]


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ValidationError("encrypted PDFs are not supported", code="encrypted_file")
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(path: Path) -> str:
    from docx import Document

    with zipfile.ZipFile(path) as archive:
        if len(archive.infolist()) > 10_000:
            raise ValidationError("abnormal Office archive", code="abnormal_archive")
        if any(item.file_size > 100 * 1024 * 1024 for item in archive.infolist()):
            raise ValidationError("abnormal Office archive", code="abnormal_archive")
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows[:200])
    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    raw = path.read_text(encoding="utf-8-sig")
    rows = csv.reader(io.StringIO(raw))
    return "\n".join(" | ".join(row[:50]) for _, row in zip(range(201), rows, strict=False))


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True, keep_links=False)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets[:5]:
            parts.append(f"[Sheet: {sheet.title}]")
            for row_index, row in enumerate(sheet.iter_rows(max_col=50, values_only=True)):
                if row_index >= 200:
                    break
                parts.append(" | ".join("" if value is None else str(value) for value in row))
    finally:
        workbook.close()
    return "\n".join(parts)


def select_attachment_context(query: str, attachments: list[tuple[str, str]]) -> str:
    keywords = {token.lower() for token in re.findall(r"[\wÀ-ỹ]{3,}", query, flags=re.UNICODE)}
    candidates: list[tuple[int, int, str]] = []
    for filename, text in attachments:
        chunks = [text[index : index + CHUNK_CHARS] for index in range(0, len(text), CHUNK_CHARS)]
        for index, chunk in enumerate(chunks):
            score = sum(chunk.lower().count(keyword) for keyword in keywords)
            boundary = 1 if index in {0, max(0, len(chunks) - 1)} else 0
            marker = f"[Nguồn: {filename} — đoạn {index + 1}]\n{chunk}"
            candidates.append((score, boundary, marker))
    candidates.sort(key=lambda item: (item[0], item[1]), reverse=True)
    selected: list[str] = []
    used = 0
    for _score, _boundary, chunk in candidates:
        if used + len(chunk) > MAX_CONTEXT_CHARS:
            continue
        selected.append(chunk)
        used += len(chunk)
        if used >= MAX_CONTEXT_CHARS:
            break
    return "\n\n".join(selected)
